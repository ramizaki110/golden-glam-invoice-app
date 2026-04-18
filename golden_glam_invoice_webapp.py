from __future__ import annotations

import base64
import json
import os
import re
import time
import urllib.error
import urllib.parse
import urllib.request
import uuid
from pathlib import Path

from flask import Flask, jsonify, request, send_from_directory
from openpyxl import load_workbook

from golden_glam_invoice_generator import draw_invoice

BASE_DIR   = Path(__file__).resolve().parent
HTML_FILE  = BASE_DIR / "GoldenGlam_InvoiceGenerator_hosted.html"
OUTPUT_DIR = BASE_DIR / "generated_invoices"
OUTPUT_DIR.mkdir(exist_ok=True)

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 32 * 1024 * 1024  # 32MB max upload

# ── Temp image hosting via Supabase Storage ───────────────────────────────────
STORAGE_BUCKET = "temp-images"
_bucket_ensured = False


def _ensure_bucket():
    global _bucket_ensured
    if _bucket_ensured or not SUPABASE_URL or not _storage_key():
        return
    try:
        key  = _storage_key()
        body = json.dumps({"id": STORAGE_BUCKET, "name": STORAGE_BUCKET, "public": True}).encode()
        req  = urllib.request.Request(
            f"{SUPABASE_URL}/storage/v1/bucket",
            data=body,
            headers={
                "apikey":        key,
                "Authorization": f"Bearer {key}",
                "Content-Type":  "application/json",
            },
            method="POST"
        )
        urllib.request.urlopen(req, timeout=8)
    except Exception as e:
        # 400 or 409 both mean bucket already exists — that's fine
        err = str(e)
        if "400" not in err and "409" not in err and "already exists" not in err.lower():
            print(f"[storage] bucket create warning: {e}")
    _bucket_ensured = True


def _storage_key():
    # Service key bypasses RLS — required for storage write operations
    return SUPABASE_SERVICE_KEY or SUPABASE_KEY


def _upload_temp_image(img_bytes: bytes):
    if not SUPABASE_URL or not _storage_key():
        print("[storage] no Supabase credentials available")
        return None, None
    _ensure_bucket()
    filename = f"tmp_{uuid.uuid4().hex}.jpg"
    try:
        key = _storage_key()
        req = urllib.request.Request(
            f"{SUPABASE_URL}/storage/v1/object/{STORAGE_BUCKET}/{filename}",
            data=img_bytes,
            headers={
                "apikey":        key,
                "Authorization": f"Bearer {key}",
                "Content-Type":  "image/jpeg",
                "cache-control": "3600",
            },
            method="POST"
        )
        urllib.request.urlopen(req, timeout=10)
        public_url = f"{SUPABASE_URL}/storage/v1/object/public/{STORAGE_BUCKET}/{filename}"
        print(f"[storage] uploaded OK: {filename}")
        return public_url, filename
    except Exception as e:
        print(f"[storage] upload FAILED: {e}")
        return None, None


def _delete_temp_image(filename: str):
    if not SUPABASE_URL or not _storage_key() or not filename:
        return
    try:
        key  = _storage_key()
        body = json.dumps({"prefixes": [filename]}).encode()
        req  = urllib.request.Request(
            f"{SUPABASE_URL}/storage/v1/object/{STORAGE_BUCKET}",
            data=body,
            headers={
                "apikey":        key,
                "Authorization": f"Bearer {key}",
                "Content-Type":  "application/json",
            },
            method="DELETE"
        )
        urllib.request.urlopen(req, timeout=8)
    except Exception as e:
        print(f"[storage] delete failed: {e}")


# ── Invoice parsing ────────────────────────────────────────────────────────────

def _clean_currency(x: str) -> float:
    return float(x.replace(",", "").strip())


def parse_summary(summary_text: str) -> dict:
    m = re.search(r"GG-INV\|([^|]+)\|([^|]+)\|ref:(.*)", summary_text)
    if not m:
        raise ValueError("Could not read invoice number/date from summary.")
    number, date, ref = m.groups()

    m = re.search(r"client:(.*?)\|no:(.*?)\|ph:(.*?)(?:\|email:(.*))?$", summary_text, re.M)
    if not m:
        raise ValueError("Could not read client section from summary.")
    client_name  = m.group(1).strip()
    client_no    = m.group(2).strip()
    client_phone = m.group(3).strip()
    client_email = (m.group(4) or "").strip()

    m = re.search(r"addr:(.+)", summary_text)
    addr_line      = m.group(1).strip() if m else ""
    client_address = [p.strip() for p in addr_line.split(",") if p.strip()]

    m = re.search(r"file:(.+)", summary_text)
    base_filename = m.group(1).strip() if m else \
        f"GG__{number}_{client_name.upper().replace(' ', '_')}"

    delivery_type = ""
    delivery_charge = 0.0
    tax_rate = 0.0
    m = re.search(
        r"\n([^\n|]+)\|sub:\$(-?[0-9.,]+)\|del:\$(-?[0-9.,]+)"
        r"\|tax\((\d+(?:\.\d+)?)%\):\$(-?[0-9.,]+)\|total:\$(-?[0-9.,]+)",
        summary_text
    )
    if m:
        delivery_type   = m.group(1).strip()
        delivery_charge = _clean_currency(m.group(3))
        tax_rate        = float(m.group(4)) / 100.0

    payment_terms = "advance"
    low = summary_text.lower()
    if "pay:in installments" in low or "pay:installments" in low \
            or "pay:payment in installments" in low:
        payment_terms = "installments"
    elif "pay:paid in advance" in low:
        payment_terms = "advance"

    notes = ""
    m = re.search(r"\nnotes:(.+?)(?:\ninstallments:|\nINTERNAL|\Z)", summary_text, re.S)
    if m:
        notes = m.group(1).strip()

    installments = []
    installment_split_type = "amount"
    m_inst = re.search(r"\ninstallments:(.+?)(?:\nINTERNAL|\nITEMS|\Z)", summary_text, re.S)
    if m_inst:
        for part in m_inst.group(1).strip().split(";"):
            seg = part.strip().split("|")
            if len(seg) >= 3:
                installments.append({"date": seg[1], "val": seg[2]})
                if len(seg) >= 4:
                    installment_split_type = seg[3]

    parts = summary_text.split("ITEMS", 1)
    if len(parts) < 2:
        raise ValueError("Could not find invoice items in summary.")
    items_section = parts[1]
    if "INTERNAL" in items_section:
        items_part, internal_part = items_section.split("INTERNAL", 1)
    else:
        items_part, internal_part = items_section, ""

    cleaned = [l for l in items_part.splitlines() if not l.strip().startswith("(fmt:")]
    items_part = "\n".join(cleaned).strip()

    item_pattern = re.compile(
        r"\[([^\]]+)\](.*?)\|([^|\n]+)\|qty:(\d+)\|\$(-?[0-9.,]+)"
        r"(?:\|disc:([0-9.]+)%?)?\|tot:\$(-?[0-9.,]+)(?:\|del:([^\n]*))?"
        r"(?:\n\s*Photo-base64:\s*(data:image\/[a-zA-Z]+;base64,[A-Za-z0-9+/=]+))?"
        r"(?=\n\[|\n[^\n]+\|sub:\$|\npay:|\nnotes:|\ninstallments:|\nINTERNAL|\Z)",
        re.S,
    )
    items = []
    for g in item_pattern.finditer(items_part):
        no, desc, unit, qty, unit_price, disc, _tot, delivery, photo = g.groups()
        items.append({
            "no":          no.strip(),
            "description": " ".join(desc.strip().split()).replace("⁄", "|"),
            "delivery":    " ".join((delivery or "").strip().split()).replace("⁄", "|"),
            "qty":         int(qty),
            "unit":        unit.strip(),
            "unit_price":  _clean_currency(unit_price),
            "discount":    (float(disc) / 100.0) if disc else 0.0,
            "image":       re.sub(r"\s+", "", photo) if photo else "",
        })
    if not items:
        raise ValueError("Could not parse any line items from summary.")

    internal_map = {}
    for line in internal_part.splitlines():
        line = line.strip()
        if not line.startswith("["):
            continue
        im = re.match(r"\[([^\]]+)\]", line)
        if not im:
            continue
        item_no = im.group(1)
        vendor_no = vendor_name = ""
        raw_cost = cost_disc = cost = 0.0
        m = re.search(r'vendor_no:\s*"([^"]*)"', line)
        if m: vendor_no = m.group(1)
        m = re.search(r'vendor:\s*"([^"]*)"', line)
        if m: vendor_name = m.group(1)
        m = re.search(r"raw_cost:\s*([0-9.]+)", line)
        if m: raw_cost = float(m.group(1))
        m = re.search(r"cost_disc:\s*([0-9.]+)%", line)
        if m: cost_disc = float(m.group(1))
        m = re.search(r"(?<![a-z_])cost:\s*([0-9.]+)", line)
        if m: cost = float(m.group(1))
        internal_map[item_no] = {
            "vendor_name": vendor_name, "vendor_no": vendor_no,
            "raw_cost": raw_cost, "cost_disc": cost_disc, "cost": cost,
        }
    for item in items:
        item.update(internal_map.get(item["no"], {}))

    return {
        "number": number.strip(), "date": date.strip(), "reference": ref.strip(),
        "client_name": client_name, "client_no": client_no,
        "client_phone": client_phone, "client_email": client_email,
        "client_address": client_address, "delivery_type": delivery_type,
        "delivery_charge": delivery_charge, "tax_rate": tax_rate,
        "payment_terms": payment_terms, "notes": notes,
        "installments": installments, "installment_split_type": installment_split_type,
        "items": items, "base_filename": base_filename,
    }


def generate_from_summary(summary_text: str) -> tuple[Path, Path]:
    invoice   = parse_summary(summary_text)
    safe_name = invoice.pop("base_filename")
    pdf_path  = OUTPUT_DIR / f"{safe_name}.pdf"
    xlsx_path = draw_invoice(invoice, str(pdf_path))
    return pdf_path, Path(xlsx_path)


# ── Supabase ───────────────────────────────────────────────────────────────────

SUPABASE_URL         = os.environ.get("SUPABASE_URL", "").rstrip("/")
SUPABASE_KEY         = os.environ.get("SUPABASE_KEY", "")
SUPABASE_SERVICE_KEY = os.environ.get("SUPABASE_SERVICE_KEY", "")
TABLE        = "gg_data"
ROW_KEY      = "main"
EMPTY_DATA   = {"clients": [], "library": [], "vendors": [], "invoices": []}


def _sb_headers():
    return {
        "apikey": SUPABASE_KEY, "Authorization": f"Bearer {SUPABASE_KEY}",
        "Content-Type": "application/json", "Prefer": "return=representation",
    }


def _load_data() -> dict:
    if not SUPABASE_URL or not SUPABASE_KEY:
        return dict(EMPTY_DATA)
    try:
        url = f"{SUPABASE_URL}/rest/v1/{TABLE}?key=eq.{ROW_KEY}&select=value"
        with urllib.request.urlopen(urllib.request.Request(url, headers=_sb_headers()), timeout=8) as r:
            rows = json.loads(r.read())
            return rows[0].get("value") or dict(EMPTY_DATA) if rows else dict(EMPTY_DATA)
    except Exception as e:
        print(f"[data] load error: {e}")
    return dict(EMPTY_DATA)


def _save_data(data: dict):
    if not SUPABASE_URL or not SUPABASE_KEY:
        return
    try:
        body = json.dumps({"key": ROW_KEY, "value": data}).encode()
        hdrs = {**_sb_headers(), "Prefer": "resolution=merge-duplicates,return=minimal"}
        req  = urllib.request.Request(f"{SUPABASE_URL}/rest/v1/{TABLE}",
                                      data=body, headers=hdrs, method="POST")
        urllib.request.urlopen(req, timeout=8)
    except Exception as e:
        print(f"[data] save error: {e}")


@app.get("/api/data")
def api_get_data():
    return jsonify(_load_data())

@app.post("/api/data")
def api_save_data():
    payload = request.get_json(silent=True) or {}
    data    = _load_data()
    for key in ("clients", "library", "vendors", "invoices"):
        if key in payload:
            data[key] = payload[key]
    _save_data(data)
    return jsonify({"ok": True})

@app.get("/api/photo/<path:key>")
def api_get_photo(key):
    if not SUPABASE_URL or not SUPABASE_KEY:
        return jsonify({"ok": False, "error": "No DB"}), 503
    try:
        url = f"{SUPABASE_URL}/rest/v1/gg_photos?key=eq.{urllib.parse.quote(key)}&select=data"
        with urllib.request.urlopen(urllib.request.Request(url, headers=_sb_headers()), timeout=8) as r:
            rows = json.loads(r.read())
            if rows:
                return jsonify({"ok": True, "data": rows[0]["data"]})
            return jsonify({"ok": False, "data": ""})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

@app.post("/api/photo")
def api_save_photo():
    if not SUPABASE_URL or not SUPABASE_KEY:
        return jsonify({"ok": False, "error": "No DB"}), 503
    payload = request.get_json(silent=True) or {}
    key  = payload.get("key", "")
    data = payload.get("data", "")
    if not key:
        return jsonify({"ok": False, "error": "Missing key"}), 400
    try:
        body = json.dumps({"key": key, "data": data}).encode()
        hdrs = {**_sb_headers(), "Prefer": "resolution=merge-duplicates,return=minimal"}
        req  = urllib.request.Request(f"{SUPABASE_URL}/rest/v1/gg_photos",
                                      data=body, headers=hdrs, method="POST")
        urllib.request.urlopen(req, timeout=10)
        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

@app.delete("/api/photo/<path:key>")
def api_delete_photo(key):
    if not SUPABASE_URL or not SUPABASE_KEY:
        return jsonify({"ok": True})
    try:
        url = f"{SUPABASE_URL}/rest/v1/gg_photos?key=eq.{urllib.parse.quote(key)}"
        urllib.request.urlopen(
            urllib.request.Request(url, headers=_sb_headers(), method="DELETE"), timeout=8)
    except Exception:
        pass
    return jsonify({"ok": True})


# ── Competitor Pricing ─────────────────────────────────────────────────────────

RENDER_URL = os.environ.get("RENDER_EXTERNAL_URL", "").rstrip("/")

# Reputable retailers — both domains and display names used in Shopping results
FURNITURE_RETAILERS = {
    # Domains
    "wayfair.com", "westelm.com", "perigold.com", "potterybarn.com",
    "crateandbarrel.com", "cb2.com", "rh.com", "restorationhardware.com",
    "serenaandlily.com", "article.com", "worldmarket.com", "target.com",
    "amazon.com", "homedepot.com", "lowes.com", "ikea.com",
    "ballarddesigns.com", "arhaus.com", "roomandboard.com", "zgallerie.com",
    "hayneedle.com", "overstock.com", "allmodern.com", "jossandmain.com",
    "anthropologie.com", "burkedecor.com", "lumens.com", "ylighting.com",
    "circafurniture.com", "furniturerow.com", "ethanallen.com",
    "livingspaces.com", "castlery.com", "tuftandneedle.com",
    "onekingslane.com", "highfashionhome.com", "jaysonhome.com",
    "luluandgeorgia.com", "topmodern.com", "luxedecor.com", "2modern.com",
    "abccarpet.com", "stashhomefurniture.com", "graysonliving.com",
    "mcgeeandco.com", "shopatironwood.com", "lauradesignco.com",
    "laylagrayce.com", "cityhome.com", "fineline.com",
    # Display names (as returned by Google Shopping source field)
    "wayfair", "west elm", "pottery barn", "crate and barrel", "cb2",
    "restoration hardware", "rh", "serena & lily", "serena and lily",
    "article", "world market", "ballard designs", "arhaus",
    "room & board", "room and board", "z gallerie", "all modern",
    "joss & main", "joss and main", "anthropologie", "burke decor",
    "lumens", "one kings lane", "high fashion home", "jayson home",
    "lulu and georgia", "topmodern", "luxedecor", "2modern",
    "abc carpet & home", "abc carpet and home", "stash home furniture",
    "grayson living", "mcgee & co", "mcgee and co", "layla grayce",
    "city home", "fine line furniture", "perigold", "perigold by wayfair", "hayneedle",
    "overstock", "target", "ikea", "amazon",
}


def _serpapi_get(params: dict, serpapi_key: str, timeout: int = 20) -> dict:
    params["api_key"] = serpapi_key
    url = "https://serpapi.com/search?" + urllib.parse.urlencode(params)
    req = urllib.request.Request(url, headers={"Accept": "application/json"})
    with urllib.request.urlopen(req, timeout=timeout) as resp:
        return json.loads(resp.read())


def _parse_price(price_str: str) -> float | None:
    """Convert '$1,099.00' or '1099' to float. Returns None if unparseable."""
    if not price_str:
        return None
    cleaned = re.sub(r"[^\d.]", "", str(price_str))
    try:
        val = float(cleaned)
        return val if 1 < val < 100_000 else None
    except ValueError:
        return None


def _shopping_results_to_rows(items: list) -> list:
    """Convert SerpAPI shopping_results into our standard row format."""
    rows = []
    for it in items:
        price = _parse_price(it.get("price", ""))
        if price is None:
            continue
        # Prefer product_link (direct URL) over link (Google redirect)
        # SerpAPI returns both — product_link goes straight to the retailer
        link = (it.get("product_link") or it.get("link") or "").strip()

        # If it's still a Google redirect URL, try to extract the actual URL
        if "google.com" in link or "ibp=oshop" in link:
            import urllib.parse as _up
            try:
                qs = _up.parse_qs(_up.urlparse(link).query)
                # Google Shopping encodes retailer URL in 'q' or 'url' param
                for param in ("url","q","adurl"):
                    if param in qs:
                        candidate = qs[param][0]
                        if candidate.startswith("http") and "google.com" not in candidate:
                            link = candidate
                            break
            except Exception:
                pass

        source   = it.get("source", "")
        domain   = urllib.parse.urlparse(link).netloc.replace("www.", "") if link else ""
        retailer = source or domain
        src_lower = source.lower()
        dom_lower = domain.lower()
        is_rep    = any(r in src_lower or r in dom_lower for r in FURNITURE_RETAILERS)
        rows.append({
            "retailer":    retailer,
            "price":       price,
            "url":         link,
            "title":       it.get("title", ""),
            "thumbnail":   it.get("thumbnail", ""),
            "reputable":   is_rep,
            "source_type": "shopping",
        })
    return rows


def _lens_results_to_rows(matches: list) -> list:
    """Convert SerpAPI Google Lens visual_matches into our standard row format."""
    rows = []
    for it in matches:
        price_raw = it.get("price")
        if isinstance(price_raw, dict):
            price = _parse_price(price_raw.get("value", ""))
        else:
            price = _parse_price(price_raw or "")
        if price is None:
            continue
        link     = it.get("link", "")
        domain   = urllib.parse.urlparse(link).netloc.replace("www.", "")
        source   = it.get("source", domain)
        rows.append({
            "retailer":    source or domain,
            "price":       price,
            "url":         link,
            "title":       it.get("title", ""),
            "thumbnail":   it.get("thumbnail", ""),
            "reputable":   any(r in domain for r in FURNITURE_RETAILERS),
            "source_type": "lens",
        })
    return rows


@app.post("/api/price-check")
def api_price_check():
    try:
        return _price_check_inner()
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({"ok": False, "error": f"Server error: {str(e)}"}), 500


def _price_check_inner():
    payload      = request.get_json(silent=True) or {}
    image_b64    = payload.get("image", "")
    product_text = payload.get("product", "").strip()
    vendor_text  = payload.get("vendor", "").strip()
    sku          = payload.get("sku", "").strip()

    SERPAPI_KEY = os.environ.get("SERPAPI_KEY", "")
    if not SERPAPI_KEY:
        return jsonify({"ok": False,
            "error": "SerpAPI not configured. Add SERPAPI_KEY to Render environment variables."}), 503

    results      = []
    product_name = product_text
    lens_name    = ""
    image_used   = False
    _lens_direct = {}  # domain -> (url, title, retailer_name) found in Lens visual matches

    # ── Step 1: Google Lens via Supabase Storage (external, no deadlock) ────────
    if image_b64:
        image_used = True

        if "," in image_b64:
            _, b64_data = image_b64.split(",", 1)
        else:
            b64_data = image_b64

        img_bytes = None
        try:
            img_bytes = base64.b64decode(b64_data)
        except Exception as e:
            print(f"[lens] base64 decode failed: {e}")

        image_url     = None
        temp_filename = None
        if img_bytes:
            image_url, temp_filename = _upload_temp_image(img_bytes)

        if image_url:
            try:
                lens_data = _serpapi_get({
                    "engine":  "google_lens",
                    "url":     image_url,
                    "hl":      "en",
                    "country": "us",
                }, SERPAPI_KEY, timeout=25)
                visual_matches = lens_data.get("visual_matches", [])
                if visual_matches:
                    import re as _re_lens

                    def _clean_title(t):
                        # Strip retailer suffixes and packaging/shipping noise
                        for pat in [
                            r'\s*[|]\s*(wayfair|pottery barn|west elm|perigold|crate.*barrel|'
                            r'restoration hardware|overstock|amazon|target|walmart|houzz|'
                            r'google shopping|bing shopping)[^|]*$',
                            r"\s*(you'?ll love|shop now|best sellers|on sale|free shipping)[^|]*$",
                            r',?\s*\d+\s*cartons?\b.*$',   # "2 Cartons - Uttermost"
                            r',?\s*set\s+of\s+\d+\b.*$',  # "Set of 2"
                            r',?\s*\d+\s*pieces?\b.*$',    # "3 Pieces"
                            r'\s*[-–]\s*[A-Z][a-z]+\s*$',  # trailing brand suffix "- Uttermost"
                        ]:
                            t = _re_lens.sub(pat, '', t, flags=_re_lens.IGNORECASE).strip()
                        return t.strip()

                    # SIMPLE approach: just use the top Lens match title, cleaned
                    # Lens already ranks by visual similarity — top result is best
                    lens_name = ""
                    for vm in visual_matches[:3]:
                        candidate = _clean_title(vm.get("title", ""))
                        if candidate and len(candidate) > 5:
                            lens_name = candidate
                            break

                    # Append detected color if not already in name
                    COLOR_WORDS = ["natural","white","black","brown","grey","gray",
                                   "beige","cream","tan","teak","vintage","walnut",
                                   "ivory","charcoal","espresso","rattan","antique"]
                    if lens_name:
                        detected_color = next(
                            (cw for vm in visual_matches[:5]
                             for cw in COLOR_WORDS
                             if cw in vm.get("title","").lower()
                             and cw not in lens_name.lower()),
                            ""
                        )
                        if detected_color:
                            lens_name = f"{lens_name} {detected_color}"

                    print(f"[lens] top_title='{visual_matches[0].get('title','')}' final_name='{lens_name}'")
                results += _lens_results_to_rows(visual_matches)
                print(f"[lens] {len(visual_matches)} visual matches")

                # Extract direct retailer URLs found by Lens (no search needed — visual match)
                DIRECT_DOMAINS = {
                    "perigold.com":    "Perigold",
                    "westelm.com":     "West Elm",
                    "potterybarn.com": "Pottery Barn",
                    "rh.com":          "RH",
                }
                lens_direct_urls = {}  # domain -> (url, title)
                for vm in visual_matches[:20]:
                    vm_url = vm.get("link","")
                    for dom, name in DIRECT_DOMAINS.items():
                        if dom in vm_url and dom not in lens_direct_urls:
                            lens_direct_urls[dom] = (vm_url, vm.get("title",""), name)
                            print(f"[lens] found direct retailer in visual: {name} -> {vm_url[:60]}")
                # Store for use in direct retailer step
                _lens_direct = lens_direct_urls
            except Exception as e:
                print(f"[lens] SerpAPI call failed: {e}")
            finally:
                _delete_temp_image(temp_filename)
        else:
            print("[lens] skipping Lens — no image URL available")

    # ── Step 2: Determine Shopping query ──────────────────────────────────────
    # Combine: vendor + product name + SKU for most precise search
    if product_text:
        base_name = product_text
    elif lens_name:
        base_name = lens_name
    elif sku:
        base_name = sku
    elif image_used:
        return jsonify({
            "ok": False,
            "error": "Google Lens could not identify this product. Please type the product name or brand in the field above.",
            "lens_failed": True,
        }), 400
    else:
        return jsonify({"ok": False,
            "error": "Please provide a product name, SKU, or upload an image."}), 400

    # Prepend vendor if not already in the base name
    if vendor_text and vendor_text.lower() not in base_name.lower():
        shopping_query = f"{vendor_text} {base_name}"
    else:
        shopping_query = base_name

    # Append SKU if provided (strong signal for exact product match)
    if sku and sku not in shopping_query:
        shopping_query = f"{shopping_query} {sku}"

    product_name = shopping_query
    print(f"[shopping] query='{shopping_query}' vendor='{vendor_text}' sku='{sku}'")


    # ── Step 3: Google Shopping — two passes ────────────────────────────────
    # Pass 1: General search
    try:
        shopping_data = _serpapi_get({
            "engine": "google_shopping",
            "q":      shopping_query,
            "gl":     "us",
            "hl":     "en",
            "num":    "20",
        }, SERPAPI_KEY, timeout=15)
        shopping_rows = _shopping_results_to_rows(shopping_data.get("shopping_results", []))
        results += shopping_rows
        print(f"[shopping] pass1 found {len(shopping_rows)} results")
    except Exception as e:
        print(f"[shopping] pass1 failed: {e}")

    # Pass 2: Targeted at reputable furniture retailers to surface Perigold,
    # West Elm, Pottery Barn etc. that may not rank organically
    # Pass 2 & 3: Targeted searches at top reputable retailers
    PREMIUM_SITES  = "site:perigold.com OR site:westelm.com OR site:rh.com OR site:potterybarn.com OR site:crateandbarrel.com OR site:cb2.com OR site:serenaandlily.com OR site:arhaus.com"
    BOUTIQUE_SITES = "site:topmodern.com OR site:2modern.com OR site:luxedecor.com OR site:luluandgeorgia.com OR site:onekingslane.com OR site:laylagrayce.com OR site:roomandboard.com OR site:burkedecor.com OR site:highfashionhome.com"

    for sites, label in [(PREMIUM_SITES, "premium"), (BOUTIQUE_SITES, "boutique")]:
        try:
            td = _serpapi_get({
                "engine": "google_shopping",
                "q":      f"{shopping_query} ({sites})",
                "gl":     "us",
                "hl":     "en",
                "num":    "10",
            }, SERPAPI_KEY, timeout=15)
            rows = _shopping_results_to_rows(td.get("shopping_results", []))
            results += rows
            print(f"[shopping] pass-{label} found {len(rows)} results")
        except Exception as e:
            print(f"[shopping] pass-{label} failed: {e}")

    # ── Step 3b: Direct site search for retailers not on Google Shopping ────────
    # West Elm, Pottery Barn, RH, Perigold don't list on Google Shopping
    import re as _re_direct

    # ── Step 3b: Direct site search — Perigold, West Elm, Pottery Barn, RH ─────
    # These retailers don't list on Google Shopping — search organically + scrape
    import re as _re_direct
    import threading

    DIRECT_RETAILERS = [
        ("Perigold",     "perigold.com"),
        ("West Elm",     "westelm.com"),
        ("Pottery Barn", "potterybarn.com"),
        ("RH",           "rh.com"),
    ]

    def _get_price_from_serpapi_result(r):
        """Extract price from SerpAPI result snippet/rich snippets."""
        # Check structured rich snippet first (most reliable)
        for rs in (r.get("rich_snippet") or {}).get("top", {}).get("extensions", []):
            pm = _re_direct.search(r'[$]([0-9][0-9,]*(?:[.][0-9]{2})?)', rs)
            if pm:
                p = _parse_price("$" + pm.group(1))
                if p: return p
        # Check snippet text
        snippet = r.get("snippet", "")
        pm = _re_direct.search(r'[$]([0-9][0-9,]*(?:[.][0-9]{2})?)', snippet)
        if pm:
            p = _parse_price("$" + pm.group(1))
            if p: return p
        return None

    def _scrape_nextjs_price(url):
        # Scrape price from Next.js / React sites that don't show price in Google snippets
        # Looks for __NEXT_DATA__ JSON blob embedded in HTML
        try:
            req = urllib.request.Request(url, headers={
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
                "Accept": "text/html,application/xhtml+xml",
                "Accept-Language": "en-US,en;q=0.9",
            })
            with urllib.request.urlopen(req, timeout=8) as resp:
                text = resp.read().decode("utf-8", errors="ignore")

            # Next.js data blob
            nd_m = _re_direct.search('id="__NEXT_DATA__"[^>]*>([{][^<]{10,})</script>', text, _re_direct.DOTALL)
            if nd_m:
                nd_str = nd_m.group(1)
                for pat in [r'"price"\s*:\s*"?(\d+(?:\.\d{2})?)"?',
                            r'"salePrice"\s*:\s*"?(\d+(?:\.\d{2})?)"?',
                            r'"currentPrice"\s*:\s*"?(\d+(?:\.\d{2})?)"?',
                            r'"listPrice"\s*:\s*"?(\d+(?:\.\d{2})?)"?']:
                    m = _re_direct.search(pat, nd_str)
                    if m:
                        v = float(m.group(1))
                        if 10 < v < 100000:
                            print(f"[scrape] Next.js price found: ${v}")
                            return v

            # JSON-LD structured data
            for ld_m in _re_direct.finditer(r'application/ld\+json[^>]*>(.*?)</script>', text, _re_direct.DOTALL):
                try:
                    m = _re_direct.search(r'"price"\s*:\s*"?(\d+(?:\.\d{2})?)"?', ld_m.group(1))
                    if m:
                        v = float(m.group(1))
                        if 10 < v < 100000:
                            return v
                except Exception:
                    pass

            # Standard HTML
            for pat in [r'itemprop="price"[^>]*content="(\d+(?:\.\d{2})?)"',
                        r'data-price="(\d+(?:\.\d{2})?)"']:
                m = _re_direct.search(pat, text)
                if m:
                    v = float(m.group(1))
                    if 10 < v < 100000:
                        return v
        except Exception as e:
            print(f"[scrape] {url[:50]}: {e}")
        return None

    direct_results = []
    already_found  = {r.get("retailer","").lower() for r in results}
    lock = threading.Lock()

    # Build tight search query using vendor + product name
    _core = product_text or lens_name or ""
    _stop_q = {"with","from","and","for","the","outdoor","indoor","occasional",
               "dining","accent","lounge","vintage","natural","white","black",
               "brown","beige","chair","table","sofa","set","patio","cartons","carton"}
    _core_words = [w.strip(".,") for w in _core.split() if w.lower() not in _stop_q and len(w) > 2]

    # Use vendor_text if provided, otherwise detect from known brands
    if vendor_text:
        _brand_phrase = vendor_text
    else:
        _known_brands = [("Four Hands","four hands"),("Pottery Barn","pottery barn"),
                         ("West Elm","west elm"),("Uttermost","uttermost"),
                         ("Arteriors","arteriors"),("Bernhardt","bernhardt"),
                         ("Restoration Hardware","restoration hardware")]
        _brand_phrase = next((b for b,bl in _known_brands if bl in _core.lower()), "")

    _product_words = [w for w in _core_words
                      if w.lower() not in _brand_phrase.lower().split()]
    _product_phrase = ' '.join(_product_words[:3])
    print(f"[direct] brand='{_brand_phrase}' product='{_product_phrase}'")

    def _is_relevant(title):
        if not _product_words:
            return True
        title_low = title.lower()
        return any(
            bool(_re_direct.search(r'\b' + _re_direct.escape(w.lower()) + r'\b', title_low))
            for w in _product_words[:2]
        )

    def _search_retailer(ret_name, domain):
        if ret_name.lower() in already_found:
            return

        # Strategy A: Use Lens visual match URL if Lens already found this retailer
        if domain in _lens_direct:
            vm_url, vm_title, _ = _lens_direct[domain]
            # Only use if the Lens title is relevant
            if _is_relevant(vm_title):
                try:
                    sr2 = _serpapi_get({
                        "engine": "google",
                        "q":      f'site:{domain} {_brand_phrase} {_product_phrase}',
                        "gl":     "us", "hl": "en", "num": "3",
                    }, SERPAPI_KEY, timeout=10)
                    for r2 in sr2.get("organic_results", []):
                        if domain not in r2.get("link",""):
                            continue
                        if not _is_relevant(r2.get("title","")):
                            continue
                        price = _get_price_from_serpapi_result(r2)
                        if price:
                            with lock:
                                direct_results.append({
                                    "retailer": ret_name, "price": price,
                                    "url": r2.get("link", vm_url),
                                    "title": r2.get("title", vm_title),
                                    "thumbnail": "", "reputable": True, "source_type": "shopping",
                                })
                            print(f"[direct] {ret_name} via Lens+search @ ${price}")
                            return
                except Exception as ex:
                    print(f"[direct] {ret_name} Lens+search failed: {ex}")
            else:
                print(f"[direct] {ret_name} Lens match irrelevant: {vm_title[:40]}")

        # Strategy B: Google search with quoted brand + product name
        # Use 3 progressively looser queries
        queries = [
            f'site:{domain} "{_brand_phrase}" "{_product_phrase}"' if _brand_phrase and _product_phrase else None,
            f'site:{domain} {_brand_phrase} {_product_phrase}' if _brand_phrase else None,
            f'site:{domain} {_product_phrase}' if _product_phrase else None,
        ]
        for q in filter(None, queries):
            try:
                sr = _serpapi_get({
                    "engine": "google",
                    "q":      q,
                    "gl":     "us", "hl": "en", "num": "5",
                }, SERPAPI_KEY, timeout=10)
                for r in sr.get("organic_results", []):
                    link = r.get("link","")
                    if domain not in link:
                        continue
                    title = r.get("title","")
                    price = _get_price_from_serpapi_result(r)
                    if not price:
                        price = _scrape_nextjs_price(link)
                    if price:
                        with lock:
                            direct_results.append({
                                "retailer": ret_name, "price": price,
                                "url": link, "title": title,
                                "thumbnail": "", "reputable": True, "source_type": "shopping",
                            })
                        print(f"[direct] {ret_name} @ ${price}: {title[:40]}")
                        return
                print(f"[direct] {ret_name}: no price in snippet for q='{q[:60]}'")
            except Exception as ex:
                print(f"[direct] {ret_name} query failed: {ex}")


    results.extend(direct_results)

    # ── Step 4: Deduplicate by URL, cap at 40 results ───────────────────────
    seen, deduped = set(), []
    for r in results:
        key = r["url"] or r["title"]
        if key and key not in seen:
            seen.add(key)
            deduped.append(r)
    if len(deduped) > 40:
        # Keep reputable retailers + first entries up to 40
        reputable = [r for r in deduped if r.get("reputable")]
        others    = [r for r in deduped if not r.get("reputable")]
        deduped   = (reputable + others)[:40]

    # ── Step 4b: Relevance scoring — token overlap between query and result title ─
    # Score each result by how many query words appear in the title (whole-word).
    # Keep all results but sort lower-scoring ones toward the bottom.
    import re as _re_rel
    _ref = (product_text or lens_name or "").lower()
    _rel_stop = {"with","from","that","this","and","for","the","home","decor",
                 "set","piece","modern","collection","a","an","of","in","at"}
    _ref_tokens = [w.strip(".,()-") for w in _ref.split()
                   if len(w) > 2 and w.strip(".,()-") not in _rel_stop]

    def _relevance_score(title: str) -> int:
        if not _ref_tokens:
            return 1
        t = title.lower()
        return sum(1 for w in _ref_tokens
                   if _re_rel.search(r'\b' + _re_rel.escape(w) + r'\b', t))

    if _ref_tokens:
        # Tag each result with a relevance score
        for r in deduped:
            r["_score"] = _relevance_score(r.get("title",""))
        # Filter out results with zero overlap (completely unrelated)
        # unless we don't have enough results
        scored = [r for r in deduped if r.get("_score",0) > 0]
        deduped = scored if len(scored) >= 5 else deduped
        print(f"[relevance] kept {len(deduped)} results (ref_tokens={_ref_tokens[:4]})")

    # ── Step 4c: Deduplicate by retailer (keep lowest price per retailer) ────
    seen_ret = {}
    for r in deduped:
        key = r.get("retailer","").lower().strip()
        if key not in seen_ret or r["price"] < seen_ret[key]["price"]:
            seen_ret[key] = r
    deduped = list(seen_ret.values())

    # ── Step 4d: Sort — Lens first, then reputable, then price ───────────────
    # Remove Walmart — not a reputable furniture source for interior design
    deduped = [r for r in deduped if "walmart" not in r.get("retailer","").lower()]
    # Sort: lens first, then reputable, then by relevance score desc, then price asc
    deduped.sort(key=lambda r: (
        r.get("source_type") != "lens",
        not r.get("reputable", False),
        -(r.get("_score", 0)),   # higher relevance score = earlier
        r["price"]
    ))
    # Clean up internal score field before returning
    for r in deduped:
        r.pop("_score", None)

    # ── Step 5: Price range ───────────────────────────────────────────────────
    if not deduped:
        return jsonify({
            "ok": True, "identified_product": product_name,
            "image_used": image_used, "results": [],
            "floor": None, "ceiling": None, "suggested": None,
            "query_used": shopping_query,
        })

    prices    = [r["price"] for r in deduped]
    floor_p   = round(min(prices), 2)
    ceiling_p = round(max(prices), 2)
    avg_p     = sum(prices) / len(prices)
    suggested = round(avg_p * 1.05 / 5) * 5   # 5% above avg, nearest $5

    return jsonify({
        "ok":                 True,
        "identified_product": product_name,
        "lens_name":          lens_name,
        "image_used":         image_used,
        "results":            deduped,
        "floor":              floor_p,
        "ceiling":            ceiling_p,
        "suggested":          suggested,
        "query_used":         shopping_query,
    })


# ── Delivery Estimate ─────────────────────────────────────────────────────────


def _fetch_page_shipping(url: str, zip_code: str, timeout: int = 8) -> str:
    # Fetch a product page and look for shipping/delivery text
    try:
        req = urllib.request.Request(url, headers={
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
            "Accept": "text/html,application/xhtml+xml",
        })
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            raw = resp.read()
            try:
                text = raw.decode("utf-8")
            except Exception:
                text = raw.decode("latin-1", errors="ignore")
        # Strip HTML tags
        import re
        plain = re.sub(r"<[^>]+>", " ", text)
        plain = re.sub(r"\s+", " ", plain)
        plain_lower = plain.lower()
        # Look for shipping/delivery snippets
        patterns = [
            r"free\s+(?:standard\s+)?(?:shipping|delivery)[^.]{0,80}",
            r"(?:white\s+glove|threshold|room\s+of\s+choice)[^.]{0,100}",
            r"(?:shipping|delivery)[^.]{0,60}?\$[\d,]+(?:\.\d{2})?",
            r"estimated\s+(?:shipping|delivery)[^.]{0,80}",
            r"ships?\s+(?:in|within|to)[^.]{0,60}",
        ]
        for pat in patterns:
            m = re.search(pat, plain_lower)
            if m:
                snippet = plain[m.start():m.end()].strip()
                if 8 < len(snippet) < 160:
                    return snippet.capitalize()
    except Exception as e:
        print(f"[delivery] fetch failed for {url[:50]}: {e}")
    return ""


@app.post("/api/delivery-estimate")
def api_delivery_estimate():
    try:
        return _delivery_estimate_inner()
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({"ok": False, "error": str(e)}), 500


def _delivery_estimate_inner():
    payload       = request.get_json(silent=True) or {}
    product_name  = payload.get("product", "").strip()
    address       = payload.get("address", "").strip()
    zip_code      = payload.get("zip", "30030").strip()
    retailer_urls = payload.get("urls", [])

    SERPAPI_KEY = os.environ.get("SERPAPI_KEY", "")
    if not SERPAPI_KEY:
        return jsonify({"ok": False, "error": "SerpAPI not configured"}), 503
    if not product_name:
        return jsonify({"ok": False, "error": "No product name provided"}), 400

    delivery_addr = address or f"zip {zip_code}"
    import re as _re
    city_state_m  = _re.search(r'([A-Za-z ]+),\s*([A-Z]{2})\s+\d{5}', delivery_addr)
    city_state    = f"{city_state_m.group(1).strip()}, {city_state_m.group(2)}" if city_state_m else zip_code

    results = []
    seen    = set()

    # Extract retailer names from the URLs passed from frontend
    retailer_names = []
    for url in retailer_urls[:8]:
        if url and "google.com" not in url:
            domain = urllib.parse.urlparse(url).netloc.replace("www.", "")
            # Convert domain to readable name
            name = domain.split(".")[0].replace("-", " ").title()
            retailer_names.append((name, domain, url))

    # Search per-retailer delivery policy — more accurate than generic product search
    for ret_name, domain, url in retailer_names[:5]:
        if domain.lower() in seen:
            continue
        query = f"{ret_name} furniture white glove delivery fee {zip_code}"
        try:
            data = _serpapi_get({
                "engine": "google",
                "q":      query,
                "gl":     "us",
                "hl":     "en",
                "num":    "3",
            }, SERPAPI_KEY, timeout=10)
            # Look for delivery cost in organic results
            shipping_info = ""
            for r in data.get("organic_results", []):
                snippet = r.get("snippet", "")
                snip_low = snippet.lower()
                if any(w in snip_low for w in ["deliver","shipping","freight","white glove"]):
                    # Extract the relevant part
                    import re as _re2
                    m = _re2.search(r"[^.]*(?:deliver|ship|freight|white glove)[^.]{0,120}", snip_low)
                    if m:
                        shipping_info = snippet[m.start():m.end()].strip().capitalize()
                        break
            if not shipping_info:
                shipping_info = "Contact retailer for delivery quote"
            results.append({
                "retailer": ret_name,
                "shipping": shipping_info,
                "url":      url,
                "reliable": "contact" not in shipping_info.lower(),
            })
            seen.add(domain.lower())
        except Exception as e:
            print(f"[delivery] {ret_name} query failed: {e}")

    # If no retailer URLs, do a general furniture delivery cost search
    if not results:
        try:
            data = _serpapi_get({
                "engine": "google_shopping",
                "q":      f"{product_name} white glove delivery {city_state}",
                "gl":     "us",
                "hl":     "en",
                "num":    "10",
            }, SERPAPI_KEY, timeout=15)
            for it in data.get("shopping_results", []):
                source   = it.get("source", "")
                if source.lower() in seen or "walmart" in source.lower():
                    continue
                delivery = it.get("delivery") or ""
                for ext in (it.get("extensions") or []):
                    if any(w in ext.lower() for w in ["white glove","threshold","freight","room of choice"]):
                        delivery = ext
                        break
                if not delivery:
                    continue
                results.append({
                    "retailer": source,
                    "shipping": delivery,
                    "url":      it.get("link",""),
                    "reliable": any(w in delivery.lower() for w in ["white glove","threshold","freight"]),
                })
                seen.add(source.lower())
        except Exception as e:
            print(f"[delivery] fallback failed: {e}")

    return jsonify({
        "ok":      True,
        "results": results[:8],
        "address": delivery_addr,
        "zip":     zip_code,
        "note":    "Furniture delivery to a residential address typically ranges $99–$399 (threshold) to $199–$499 (white glove with assembly). Verify exact cost with each retailer.",
    })


# ── Core routes ────────────────────────────────────────────────────────────────

@app.get("/")
def home():
    return send_from_directory(BASE_DIR, HTML_FILE.name)

@app.post("/generate")
def generate():
    payload = request.get_json(silent=True) or {}
    summary = (payload.get("summary") or "").strip()
    if not summary:
        return jsonify({"ok": False, "error": "Missing invoice summary."}), 400
    try:
        pdf_path, xlsx_path = generate_from_summary(summary)
        return jsonify({
            "ok": True,
            "pdf_name":  pdf_path.name,  "xlsx_name": xlsx_path.name,
            "pdf_url":   f"/download/{pdf_path.name}",
            "xlsx_url":  f"/download/{xlsx_path.name}",
        })
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

@app.get("/download/<path:filename>")
def download(filename: str):
    return send_from_directory(OUTPUT_DIR, filename, as_attachment=True)

@app.get("/download/<path:filename>")
def download(filename: str):
    return send_from_directory(OUTPUT_DIR, filename, as_attachment=True)


# ── Proposal Generator ─────────────────────────────────────────────────────────
# Strategy: all boilerplate (intro, principles, what we'll do, sign-off) is
# hardcoded — zero tokens spent on it. Claude ONLY writes the room descriptions
# (3-5 sentences each) and intro paragraphs. Uses claude-haiku-4-5 (~20x cheaper
# than Opus) since it only needs to write polished short paragraphs.

ANTHROPIC_KEY = os.environ.get("ANTHROPIC_API_KEY", "")

# Compact system prompt — no examples, no lengthy instructions
_PROPOSAL_SYS = (
    "You are Rana Salah, co-founder of Golden Glam Interiors LLC, an interior styling "
    "business in Atlanta. Write warm, specific, professional interior design proposal copy. "
    "Style: British English spelling, no em dashes, no AI filler phrases, no vague promises. "
    "Be concrete about what will actually change in each room. Neutral/minimalist aesthetic. "
    "3-5 sentences per room description."
)

# Only ask Claude for the variable parts
_PROPOSAL_USER = """Write proposal copy for client: {client_name}
Preferences: {preferences}

Return ONLY valid JSON, no markdown:
{{
  "intro_para_1": "2-sentence thank-you opening personalised to the client",
  "intro_para_2": "2-sentence paragraph acknowledging their specific preferences/home",
  "phases": [{phase_items}]
}}

For each phase item use:
{{"phase_number":N,"phase_name":"...","rooms":[{{"label":"(a) Name","description":"3-5 sentence specific scope"}}]}}

Phases and rooms to write:
{phases_text}"""


def _call_haiku(system: str, user: str) -> str:
    """Call claude-haiku — cheapest model, fast, good enough for proposal copy."""
    if not ANTHROPIC_KEY:
        raise RuntimeError("ANTHROPIC_API_KEY not set in environment.")
    body = json.dumps({
        "model": "claude-haiku-4-5-20251001",
        "max_tokens": 2500,
        "system": system,
        "messages": [{"role": "user", "content": user}]
    }).encode()
    req = urllib.request.Request(
        "https://api.anthropic.com/v1/messages",
        data=body,
        headers={
            "x-api-key":          ANTHROPIC_KEY,
            "anthropic-version":  "2023-06-01",
            "content-type":       "application/json",
        },
        method="POST"
    )
    try:
        with urllib.request.urlopen(req, timeout=60) as resp:
            data = json.loads(resp.read())
            return data["content"][0]["text"]
    except urllib.error.HTTPError as e:
        raise RuntimeError(f"Anthropic API error {e.code}: {e.read().decode()[:400]}")


def _build_proposal_docx(client_name: str, proposal_data: dict) -> Path:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Page margins
    for sec in doc.sections:
        sec.page_width    = int(8.5 * 914400)
        sec.page_height   = int(11  * 914400)
        sec.left_margin   = int(1.0 * 914400)
        sec.right_margin  = int(1.0 * 914400)
        sec.top_margin    = int(1.0 * 914400)
        sec.bottom_margin = int(1.0 * 914400)

    def r(run, sz=11, bold=False, color=None):
        run.font.name = "Calibri"
        run.font.size = Pt(sz)
        run.font.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)

    def para(text="", sz=11, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.LEFT,
             sb=5, sa=5, underline=False):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(sb)
        p.paragraph_format.space_after  = Pt(sa)
        if text:
            run = p.add_run(text)
            r(run, sz, bold, color)
            run.font.underline = underline
        return p

    def bul(bold_part="", rest=""):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        if bold_part:
            r1 = p.add_run(bold_part); r(r1, 11, True)
            r2 = p.add_run(rest);      r(r2, 11)
        else:
            rx = p.add_run(rest or bold_part); r(rx, 11)

    def num(bold_part="", rest=""):
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        r1 = p.add_run(bold_part); r(r1, 11, True)
        if rest:
            r2 = p.add_run(rest); r(r2, 11)

    def sub(text):
        p = doc.add_paragraph(style="List Bullet 2")
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        rx = p.add_run(text); r(rx, 10)

    def heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
        rx = p.add_run(text); r(rx, 11, True)
        rx.font.underline = True

    # ── Header (logo) ─────────────────────────────────────────────────────────
    logo_path = BASE_DIR / "golden_glam_logo_final.png"
    if logo_path.exists():
        hdr = doc.sections[0].header
        hp  = hdr.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hp.add_run().add_picture(str(logo_path), width=Inches(2.6))

    # ── Footer ────────────────────────────────────────────────────────────────
    ftr = doc.sections[0].footer
    for fp in ftr.paragraphs:
        for fx in fp.runs: fx.text = ""

    def fline(text, bold=False, sz=8):
        p = ftr.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        rx = p.add_run(text); r(rx, sz, bold, (120, 120, 120))

    fline("GOLDEN GLAM INTERIORS LLC", bold=True, sz=9)
    fline("Address: 828 Highland Ln Ne, Apt. 2204, Atlanta, GA 30306  |  Phone: 770-375-7343")
    fline("Bank account #: 930283558  Routing number: 061092387  |  Zelle email: rana_salah@goldenglam.nl")
    fline("E-mail: sales@goldenglam.nl  |  Instagram: www.instagram/goldenglam.nl  |  www.goldenglam.nl")
    pgp = ftr.add_paragraph()
    pgp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pgp.paragraph_format.space_before = Pt(2)
    rx = pgp.add_run("Page "); r(rx, 8, color=(120, 120, 120))
    for tag, txt in [("w:fldChar", None), ("w:instrText", "PAGE"), ("w:fldChar", None)]:
        el = OxmlElement(tag)
        if txt:
            el.set(qn("xml:space"), "preserve"); el.text = txt
        else:
            el.set(qn("w:fldCharType"), "begin" if not txt else "end")
        pgp._p.append(el)
    end_el = OxmlElement("w:fldChar"); end_el.set(qn("w:fldCharType"), "end")
    pgp._p.append(end_el)

    # ── Greeting ──────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(10)
    rx = p.add_run(f"Dear {client_name.strip()},"); r(rx, 13, True)

    for key in ["intro_para_1", "intro_para_2"]:
        txt = proposal_data.get(key, "")
        if txt: para(txt, sb=4, sa=5)

    # ── Boilerplate sections (zero AI tokens) ─────────────────────────────────
    heading("Design Principles")
    para("Our approach to styling and designing your home revolves around three core principles:", sb=3, sa=3)
    bul("Cohesion and unity: ", "Establishing a consistent theme, colour palette, and design style to create harmony throughout your home (from room to room)")
    bul("Balance: ", "Achieving visual harmony by carefully balancing furniture, colours, patterns, and textures - through (a)symmetrical arrangements, depending on the desired aesthetic")
    bul("Functionality and practicality: ", "Designing spaces that are intuitive, purposeful, and meet your daily needs while enhancing flow and comfort")

    heading("What We'll Do")
    num("Optimise and visualise your layout")
    sub("Conduct site measurements")
    sub("Develop detailed space plans and furniture placements")
    sub("Provide 3D furniture layout renderings to optimise space, flow and functionality")
    sub("Offer expert design insights to help you visualise the transformed space")
    num("Style and personalise your home")
    sub("Collaborate through an interactive process to refine styling concepts")
    sub("Decide on current 'to stay' and new furniture purchases")
    sub("Create personalised mood board including furniture pieces, colour palettes, materials, and textiles")
    sub("Select decorative elements such as wallpaper, wall paint, lighting, rugs, and accessories")
    num("Dress your space")
    sub("Present furniture options with details on size, style, and pricing")
    sub("Test functionality with onsite furniture mockups before purchase")
    sub("Deliver uniquely curated furniture pieces to your doorstep, thoughtfully chosen to reflect your style and personality")
    sub("Style and arrange furniture onsite ensuring a cohesive and polished look")
    num("Project management")
    sub("Guide contractor work, including guidance on sizing, painting (colour palettes), and placement")

    para("As part of styling, we will:", sb=8, sa=3)
    bul("Firstly, ",  "create a focal point to anchor each space and design around it")
    bul("Secondly, ", "curate furniture layouts to offer functionality and beauty")
    bul("Thirdly, ",  "enhance lighting through layering and ambiance")
    bul("Fourthly, ", "select furniture options, complementary wall treatments, textiles, and window treatments to complement the furniture and mood of each room")
    bul("Fifthly, ",  "add finishing touches with wall decor, art, plants, and accessories to personalise and bring your vision to life")

    # ── Phase summary table ───────────────────────────────────────────────────
    phases = proposal_data.get("phases", [])
    para("We have taken a phased approach to style and furnish your home:", sb=10, sa=5)
    tbl = doc.add_table(rows=1 + len(phases), cols=3)
    tbl.style = "Table Grid"
    tbl.columns[0].width = Inches(0.6)
    tbl.columns[1].width = Inches(5.5)
    tbl.columns[2].width = Inches(1.4)
    hdr_cells = tbl.rows[0].cells
    for cell, txt in zip(hdr_cells, ["Phase", "Interior Styling Scope", "Design Fee"]):
        cell.text = txt
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if cell.paragraphs[0].runs:
            r(cell.paragraphs[0].runs[0], 10, True, (255,255,255))
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),"1A1A1A")
        tcPr.append(shd)
    for i, ph in enumerate(phases):
        row  = tbl.rows[i+1].cells
        fill = "F5F5F5" if i % 2 == 0 else "EBEBEB"
        rooms_txt = ", ".join(rm["label"].lstrip("(abcdefghij) ") for rm in ph.get("rooms",[]))
        row[0].text = str(ph.get("phase_number", i+1))
        row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[1].text = f"{ph.get('phase_name','')}  -  {rooms_txt}"
        row[2].text = ph.get("phase_price","")
        row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for cell in row:
            if cell.paragraphs[0].runs:
                r(cell.paragraphs[0].runs[0], 10)
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"), fill)
            tcPr.append(shd)

    # ── Phase detail pages ────────────────────────────────────────────────────
    for ph in phases:
        doc.add_page_break()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(6)
        rx = p.add_run(f"Phase {ph.get('phase_number','')}: {ph.get('phase_name','')}"); r(rx, 13, True)
        rx.font.underline = True
        heading("Scope:")
        para("This space includes interior styling and decorating of:", sb=2, sa=5)
        for room in ph.get("rooms", []):
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_before = Pt(7)
            p2.paragraph_format.space_after  = Pt(2)
            rx2 = p2.add_run(room.get("label","")); r(rx2, 11, True)
            para(room.get("description",""), sb=2, sa=5)
        heading("Price:")
        price_txt = f"The interior styling and advice package for Phase {ph.get('phase_number','')}: {ph.get('phase_name','')} will be {ph.get('phase_price','')}"
        para(price_txt, sb=2, sa=8)

    # ── Sign-off ──────────────────────────────────────────────────────────────
    doc.add_paragraph()
    para("We are truly excited to bring your dream home to life and look forward to collaborating with you on this project.", sb=8, sa=6)
    para("Rana Salah", bold=True, sb=4, sa=2)
    para("Co-founder of Golden Glam Interiors LLC", sb=0, sa=4)

    safe = "".join(c if c.isalnum() or c in " _-" else "_" for c in client_name).strip().replace(" ","_")
    out  = OUTPUT_DIR / f"GoldenGlam_Proposal_{safe}.docx"
    doc.save(str(out))
    return out


@app.post("/proposal/generate")
def proposal_generate():
    payload     = request.get_json(silent=True) or {}
    client_name = (payload.get("client_name") or "").strip()
    phases_raw  = payload.get("phases") or []
    preferences = (payload.get("preferences") or "").strip()

    if not client_name:
        return jsonify({"ok": False, "error": "Client name is required."}), 400
    if not phases_raw:
        return jsonify({"ok": False, "error": "At least one phase is required."}), 400

    # Build a compact phases text for the prompt
    lines = []
    for i, ph in enumerate(phases_raw, 1):
        lines.append(f"Phase {i} - {ph.get('name','')}: {ph.get('price','')}")
        for rm in ph.get("rooms", []):
            lines.append(f"  {rm.get('label','')}: {rm.get('notes','')}")
    phases_text = "\n".join(lines)

    # Build phase_items skeleton for JSON structure guidance
    phase_items = ", ".join(
        '{{"phase_number":{i},"phase_name":"{n}","rooms":[{rooms}]}}'.format(
            i=i+1,
            n=ph.get("name",""),
            rooms=", ".join('{{"label":"{l}","description":"..."}}'.format(l=rm.get("label","")) for rm in ph.get("rooms",[]))
        )
        for i, ph in enumerate(phases_raw)
    )

    prompt = _PROPOSAL_USER.format(
        client_name=client_name,
        preferences=preferences or "None specified.",
        phase_items=phase_items,
        phases_text=phases_text
    )

    try:
        raw  = _call_haiku(_PROPOSAL_SYS, prompt)
        raw  = re.sub(r"^```[a-z]*\n?", "", raw.strip())
        raw  = re.sub(r"\n?```$", "", raw.strip())
        data = json.loads(raw)
        # Inject prices from user input (not from AI)
        for i, ph in enumerate(data.get("phases", [])):
            if i < len(phases_raw):
                ph["phase_price"] = phases_raw[i].get("price", "TBD")
        docx_path = _build_proposal_docx(client_name, data)
        return jsonify({"ok": True, "docx_name": docx_path.name, "docx_url": f"/download/{docx_path.name}"})
    except json.JSONDecodeError as e:
        return jsonify({"ok": False, "error": f"AI returned invalid JSON: {e}. Raw: {raw[:300]}"}), 500
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


if __name__ == "__main__":
    port = int(os.environ.get("PORT", "5000"))
    app.run(host="0.0.0.0", port=port, debug=False)
